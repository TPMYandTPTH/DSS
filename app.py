import os
import zipfile
from flask import Flask, render_template, request, send_file, redirect, url_for, flash
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import tempfile
import shutil

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def split_docx(input_path, output_dir, logo_path=None, font_path=None):
    """Split a Word document into 2-page chunks with company branding"""
    doc = Document(input_path)
    
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Get total sections (approximate pages)
    sections = []
    current_section = []
    
    # Simple approach: split by paragraphs (approximate)
    # This is a simplified approach - you might need to adjust based on your document structure
    paragraphs_per_page = 25  # Adjust based on your document
    paragraphs_per_job = paragraphs_per_page * 2
    
    for i, para in enumerate(doc.paragraphs):
        current_section.append(para)
        
        # When we have enough paragraphs for a job, create a new document
        if len(current_section) >= paragraphs_per_job or i == len(doc.paragraphs) - 1:
            sections.append(current_section)
            current_section = []
    
    # Create documents for each section
    output_files = []
    
    for i, section in enumerate(sections):
        new_doc = Document()
        
        # Add company logo to header if available
        if logo_path and os.path.exists(logo_path):
            section = new_doc.sections[0]
            header = section.header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = header_para.add_run()
            run.add_picture(logo_path, width=Inches(1.0))
        
        # Copy paragraphs
        for para in section:
            new_para = new_doc.add_paragraph()
            new_para.text = para.text
            
            # Apply company font if available
            if font_path and os.path.exists(font_path):
                # Note: Applying custom fonts in python-docx is complex
                # This is a simplified approach
                try:
                    new_para.style.font.name = 'Company Font'
                except:
                    pass  # Fallback to default if font can't be applied
        
        # Save the document
        output_path = os.path.join(output_dir, f'job_description_{i+1}.docx')
        new_doc.save(output_path)
        output_files.append(output_path)
    
    return output_files

def create_zip(files, zip_path):
    """Create a zip file containing all the split documents"""
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for file in files:
            zipf.write(file, os.path.basename(file))
    return zip_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        flash('No file part')
        return redirect(request.url)
    
    file = request.files['file']
    
    if file.filename == '':
        flash('No selected file')
        return redirect(request.url)
    
    if file and file.filename.endswith('.docx'):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Create a temporary directory for output files
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Get paths to company assets
            logo_path = os.path.join('assets', 'company_logo.png')
            font_path = os.path.join('assets', 'company_font.ttf')
            
            # Split the document
            output_files = split_docx(filepath, temp_dir, logo_path, font_path)
            
            # Create a zip file with all outputs
            zip_path = os.path.join(temp_dir, 'split_documents.zip')
            create_zip(output_files, zip_path)
            
            # Store the zip path in session for download
            session['zip_path'] = zip_path
            
            return render_template('result.html', count=len(output_files))
        
        finally:
            # Clean up the uploaded file
            os.remove(filepath)
    
    flash('Please upload a valid Word document (.docx)')
    return redirect(url_for('index'))

@app.route('/download')
def download_file():
    zip_path = session.get('zip_path')
    
    if zip_path and os.path.exists(zip_path):
        return send_file(
            zip_path,
            as_attachment=True,
            download_name='split_documents.zip',
            mimetype='application/zip'
        )
    
    flash('No file available for download')
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)
